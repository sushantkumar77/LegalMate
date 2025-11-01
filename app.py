import streamlit as st
import re
import io
import google.generativeai as genai
from docx import Document
import os
from datetime import datetime
from dotenv import load_dotenv

load_dotenv()

st.set_page_config(
    page_title="Legal Document Processor",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
    .main {
        padding: 0rem 1rem;
    }
    .stButton>button {
        width: 100%;
        border-radius: 10px;
        height: 3em;
        font-weight: 600;
        background-color: #007bff;
        color: white;
    }
    .upload-section {
        border: 2px dashed #4CAF50;
        border-radius: 10px;
        padding: 2rem;
        text-align: center;
        background-color: #f0f8ff;
    }
    .upload-section h2 {
        color: #1a1a1a;
        margin-bottom: 0.5rem;
    }
    .upload-section p {
        color: #333333;
    }
    .chat-message {
        padding: 1rem;
        border-radius: 10px;
        margin-bottom: 1rem;
    }
    .user-message {
        background-color: #007bff;
        color: white;
        margin-left: 20%;
    }
    .assistant-message {
        background-color: #e9ecef;
        color: #212529;
        margin-right: 20%;
    }
    .placeholder-box {
        padding: 0.8rem;
        border-radius: 8px;
        margin-bottom: 0.5rem;
        border-left: 4px solid;
        color: #212529;
    }
    .completed {
        background-color: #d4edda;
        border-color: #28a745;
        color: #155724;
    }
    .current {
        background-color: #cce5ff;
        border-color: #007bff;
        color: #004085;
    }
    .pending {
        background-color: #f8f9fa;
        border-color: #dee2e6;
        color: #6c757d;
    }
    .success-box {
        padding: 1.5rem;
        border-radius: 10px;
        background-color: #d4edda;
        border: 2px solid #28a745;
        text-align: center;
    }
    .success-box h1 {
        color: #155724;
    }
    .success-box p {
        color: #155724;
    }
    .doc-preview {
        background-color: #ffffff;
        padding: 1.5rem;
        border-radius: 10px;
        border: 2px solid #dee2e6;
        max-height: 500px;
        overflow-y: auto;
        font-family: monospace;
        white-space: pre-wrap;
        color: #212529;
    }
</style>
""", unsafe_allow_html=True)

if 'step' not in st.session_state:
    st.session_state.step = 'upload'
if 'placeholders' not in st.session_state:
    st.session_state.placeholders = []
if 'filled_data' not in st.session_state:
    st.session_state.filled_data = {}
if 'messages' not in st.session_state:
    st.session_state.messages = []
if 'current_index' not in st.session_state:
    st.session_state.current_index = 0
if 'document_text' not in st.session_state:
    st.session_state.document_text = ""
if 'file_name' not in st.session_state:
    st.session_state.file_name = ""
if 'completed_doc' not in st.session_state:
    st.session_state.completed_doc = ""
if 'api_configured' not in st.session_state:
    st.session_state.api_configured = False

GEMINI_API_KEY = os.getenv('GEMINI_API_KEY', '')

if GEMINI_API_KEY and not st.session_state.api_configured:
    try:
        genai.configure(api_key=GEMINI_API_KEY)
        model = genai.GenerativeModel('gemini-pro')
        st.session_state.api_configured = True
    except Exception as e:
        st.error("Failed to configure Gemini API. Please check your API key in .env file.")

def parse_docx(file):
    try:
        doc = Document(file)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        st.error(f"Error parsing DOCX: {str(e)}")
        return ""

def detect_placeholders_with_ai(text):
    if not st.session_state.api_configured:
        st.error("❌ Gemini API is not configured. Please add GEMINI_API_KEY to your .env file.")
        return []
    
    try:
        model = genai.GenerativeModel('gemini-pro')
        
        prompt = f"""Analyze this legal document and identify ALL placeholders that need to be filled in.
        
Document:
{text}

Instructions:
1. Find ALL placeholders in formats like: [Field], {{Field}}, <Field>, ${{Field}}, [[Field]], _____, ...
2. Also identify any fields that appear to be blank or need filling even if not in standard format
3. Return ONLY a JSON array of objects with this exact structure:
{{"placeholders": [{{"label": "exact placeholder text", "original": "original text with brackets", "position": character_position}}]}}

Be thorough and find every single placeholder that needs filling. Include the exact original text as it appears in the document.
Return valid JSON only, no explanations."""

        response = model.generate_content(prompt)
        ai_text = response.text.strip()
        
        # Clean up response to extract JSON
        if '```json' in ai_text:
            ai_text = ai_text.split('```json')[1].split('```')[0].strip()
        elif '```' in ai_text:
            ai_text = ai_text.split('```')[1].split('```')[0].strip()
        
        import json
        result = json.loads(ai_text)
        
        placeholders = []
        for idx, item in enumerate(result.get('placeholders', [])):
            key = re.sub(r'[^a-z0-9]', '_', item['label'].lower())
            placeholders.append({
                'key': key,
                'label': item['label'],
                'original': item.get('original', item['label']),
                'value': '',
                'position': item.get('position', idx),
                'type': 'ai_detected'
            })
        
        return sorted(placeholders, key=lambda x: x['position'])
        
    except Exception as e:
        st.error(f"AI placeholder detection failed: {str(e)}")
        # Fallback to regex detection
        return detect_placeholders_regex(text)

def detect_placeholders_regex(text):
    patterns = [
        (r'\[([^\]]+)\]', 'square'),
        (r'\{([^}]+)\}', 'curly'),
        (r'\$\{([^}]+)\}', 'dollar'),
        (r'<([^>]+)>', 'angle'),
        (r'\[\[([^\]]+)\]\]', 'double-square'),
        (r'\{\{([^}]+)\}\}', 'double-curly'),
        (r'_{3,}', 'underscore'),
        (r'\.{3,}', 'dots'),
    ]
    
    found = {}
    locations = []
    
    for pattern, pattern_type in patterns:
        for match in re.finditer(pattern, text):
            if match.groups():
                placeholder = match.group(1).strip()
            else:
                placeholder = match.group(0).strip()
            
            if len(placeholder) < 2 or len(placeholder) > 100:
                continue
            if placeholder.upper() in ['THE', 'AND', 'OR', 'OF', 'IN', 'TO', 'A', 'IS', 'IT', 'AT']:
                continue
            
            key = re.sub(r'[^a-z0-9]', '_', placeholder.lower())
            
            if key not in found:
                found[key] = True
                locations.append({
                    'key': key,
                    'label': placeholder,
                    'original': match.group(0),
                    'value': '',
                    'position': match.start(),
                    'type': pattern_type
                })
    
    return sorted(locations, key=lambda x: x['position'])

def get_ai_question_for_placeholder(placeholder, filled_data, document_context=""):
    if not st.session_state.api_configured:
        return f"What is the {placeholder['label']}?"
    
    try:
        model = genai.GenerativeModel('gemini-pro')
        
        # Get surrounding context from document
        doc_text = st.session_state.document_text
        placeholder_pos = doc_text.find(placeholder['original'])
        context_start = max(0, placeholder_pos - 200)
        context_end = min(len(doc_text), placeholder_pos + 200)
        surrounding_text = doc_text[context_start:context_end]
        
        prompt = f"""You are an AI assistant helping someone fill out a legal document (appears to be a SAFE agreement). 

Current placeholder to ask about: "{placeholder['label']}" (appears as "{placeholder['original']}" in document)

Document context around this placeholder:
"...{surrounding_text}..."

Already filled information:
{chr(10).join([f"- {p['label']}: {filled_data.get(p['key'], 'Not filled')}" for p in st.session_state.placeholders[:st.session_state.current_index]]) if filled_data else "This is the first field"}

Your task: Generate a natural, conversational question that:
1. Clearly explains what information is needed
2. Provides helpful context if the placeholder name is unclear
3. Suggests format or examples if appropriate (e.g., dates, amounts, jurisdictions)
4. Is professional but friendly

Examples of good questions:
- "What is the company's full legal name?" (for Company Name)
- "What amount is the investor contributing? (e.g., $100,000)" (for Purchase Amount)
- "What date is this SAFE agreement being signed? (e.g., January 15, 2024)" (for Date)
- "In which state is the company incorporated? (e.g., Delaware, California)" (for State of Incorporation)

Return ONLY the question, nothing else. Keep it under 25 words."""

        response = model.generate_content(prompt)
        question = response.text.strip()
        
        # Remove quotes if AI added them
        question = question.strip('"\'').strip()
        
        # Ensure it ends with a question mark
        if not question.endswith('?'):
            question += '?'
        
        return question
        
    except Exception as e:
        st.error(f"Error generating question: {str(e)}")
        return f"What is the {placeholder['label']}?"

def validate_response_with_ai(user_message, current_placeholder, filled_data):
    if not st.session_state.api_configured:
        return {
            'response': f"Recorded: {user_message}",
            'should_proceed': True,
            'validated_value': user_message
        }
    
    try:
        model = genai.GenerativeModel('gemini-pro')
        
        # Determine what type of field this is
        field_type = current_placeholder['label'].lower()
        
        prompt = f"""You are an AI assistant validating user input for a legal document (SAFE agreement).

Field being filled: "{current_placeholder['label']}"
User's response: "{user_message}"

Previous fields filled:
{chr(10).join([f"- {p['label']}: {filled_data.get(p['key'], 'N/A')}" for p in st.session_state.placeholders[:st.session_state.current_index]]) if filled_data else "This is the first field"}

Your tasks:
1. Determine if "{user_message}" is a valid, appropriate response for "{current_placeholder['label']}"
2. Check for common issues:
   - For amounts: Should include currency and numbers (e.g., $50,000)
   - For dates: Should be a clear date format
   - For names: Should be properly capitalized and complete
   - For states/jurisdictions: Should be valid state/country names
   - For unclear placeholders (like underscores): Ask what information is needed
3. If valid: Give a brief, encouraging acknowledgment (10-15 words max)
4. If invalid/unclear: Politely explain the issue and ask for correction (20-25 words max)

Response format:
- If VALID: "Great! [brief acknowledgment]" or "Perfect, [brief comment]"
- If INVALID: "I need [explain what's wrong]. Could you provide [what you need]?"

Be conversational and helpful. Return ONLY your response."""

        response = model.generate_content(prompt)
        ai_text = response.text.strip().strip('"\'')
        
        # Determine if we should proceed based on AI response
        proceed_indicators = ['great', 'perfect', 'excellent', 'good', 'thank', 'recorded', 'got it']
        stop_indicators = ['could you', 'please provide', 'i need', 'can you', 'clarify', 'invalid', 'must be', 'should be']
        
        ai_lower = ai_text.lower()
        should_proceed = any(indicator in ai_lower for indicator in proceed_indicators) and \
                        not any(indicator in ai_lower for indicator in stop_indicators)
        
        return {
            'response': ai_text,
            'should_proceed': should_proceed,
            'validated_value': user_message if should_proceed else None
        }
    except Exception as e:
        st.error(f"Validation error: {str(e)}")
        return {
            'response': f"Recorded: {user_message}",
            'should_proceed': True,
            'validated_value': user_message
        }

def generate_completed_document_with_ai():
    if not st.session_state.api_configured:
        return generate_completed_document_simple()
    
    try:
        model = genai.GenerativeModel('gemini-pro')
        
        replacements = "\n".join([f"Replace '{p['original']}' with '{st.session_state.filled_data.get(p['key'], p['original'])}'" 
                                  for p in st.session_state.placeholders])
        
        prompt = f"""Replace all placeholders in this legal document with the provided values. Maintain exact formatting and structure.

Original Document:
{st.session_state.document_text}

Replacements:
{replacements}

Return ONLY the completed document with placeholders replaced. Do not add explanations or modifications."""

        response = model.generate_content(prompt)
        completed = response.text.strip()
        
        # Clean up if AI added markdown code blocks
        if '```' in completed:
            completed = completed.split('```')[1]
            if completed.startswith('text') or completed.startswith('plaintext'):
                completed = '\n'.join(completed.split('\n')[1:])
            completed = completed.strip()
        
        return completed
        
    except Exception as e:
        st.warning(f"AI generation failed, using fallback method: {str(e)}")
        return generate_completed_document_simple()

def generate_completed_document_simple():
    completed = st.session_state.document_text
    for placeholder in st.session_state.placeholders:
        value = st.session_state.filled_data.get(placeholder['key'], placeholder['original'])
        completed = completed.replace(placeholder['original'], value)
    return completed

def reset_app():
    st.session_state.step = 'upload'
    st.session_state.placeholders = []
    st.session_state.filled_data = {}
    st.session_state.messages = []
    st.session_state.current_index = 0
    st.session_state.document_text = ""
    st.session_state.file_name = ""
    st.session_state.completed_doc = ""
    st.rerun()

# Check API configuration
if not GEMINI_API_KEY:
    st.error("⚠️ Gemini API Key not found! Please add GEMINI_API_KEY to your .env file.")
    st.info("Get your API key from: https://makersuite.google.com/app/apikey")
    st.stop()

col1, col2 = st.columns([6, 1])
with col1:
    st.title("📄 Legal Document Processor")
    st.caption("AI-powered document completion with intelligent placeholder detection")
with col2:
    if st.session_state.step != 'upload':
        if st.button("🔄 Reset", use_container_width=True):
            reset_app()

if st.session_state.step == 'upload':
    st.markdown("---")
    
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.markdown("""
        <div class="upload-section">
            <h2>📤 Upload Your Document</h2>
            <p>Upload a legal document with placeholders and AI will help you fill them in</p>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("<br>", unsafe_allow_html=True)
        
        uploaded_file = st.file_uploader(
            "Choose a file",
            type=['txt', 'docx'],
            help="Supports .txt and .docx files up to 10MB"
        )
        
        if uploaded_file:
            with st.spinner("🤖 AI is analyzing your document..."):
                if uploaded_file.name.endswith('.docx'):
                    text = parse_docx(uploaded_file)
                else:
                    text = uploaded_file.read().decode('utf-8')
                
                if text:
                    st.session_state.document_text = text
                    st.session_state.file_name = uploaded_file.name
                    
                    placeholders = detect_placeholders_with_ai(text)
                    
                    if placeholders:
                        st.session_state.placeholders = placeholders
                        
                        first_question = get_ai_question_for_placeholder(
                            placeholders[0], 
                            {}, 
                            text[:500]  # Pass document context
                        )
                        
                        st.session_state.messages = [{
                            'type': 'assistant',
                            'content': f"I've analyzed your document and found {len(placeholders)} placeholders to fill. Let's start! {first_question}",
                            'timestamp': datetime.now()
                        }]
                        st.session_state.step = 'chat'
                        st.rerun()
                    else:
                        st.error("❌ No placeholders detected in your document. Please upload a document with placeholders.")
        
        st.markdown("<br>", unsafe_allow_html=True)
        
        with st.expander("ℹ️ AI-Powered Detection"):
            st.markdown("""
            The AI automatically detects:
            - Standard formats: [Field], {Field}, <Field>, etc.
            - Custom placeholders and blank spaces
            - Context-aware field identification
            - Smart validation of your inputs
            """)

elif st.session_state.step == 'chat':
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.markdown("### 💬 Fill in the Details")
        
        progress = st.session_state.current_index / len(st.session_state.placeholders)
        st.progress(progress)
        st.caption(f"Progress: {st.session_state.current_index} of {len(st.session_state.placeholders)} completed")
        
        st.markdown("---")
        
        chat_container = st.container()
        with chat_container:
            for msg in st.session_state.messages:
                if msg['type'] == 'user':
                    st.markdown(f"""
                    <div class="chat-message user-message">
                        <strong>You:</strong> {msg['content']}
                    </div>
                    """, unsafe_allow_html=True)
                else:
                    st.markdown(f"""
                    <div class="chat-message assistant-message">
                        <strong>AI Assistant:</strong> {msg['content']}
                    </div>
                    """, unsafe_allow_html=True)
        
        st.markdown("---")
        
        if st.session_state.current_index < len(st.session_state.placeholders):
            with st.form(key='chat_form', clear_on_submit=True):
                user_input = st.text_input(
                    "Your answer:",
                    placeholder="Type your response here...",
                    key='user_input'
                )
                submit = st.form_submit_button("Send", use_container_width=True)
                
                if submit and user_input:
                    st.session_state.messages.append({
                        'type': 'user',
                        'content': user_input,
                        'timestamp': datetime.now()
                    })
                    
                    current_placeholder = st.session_state.placeholders[st.session_state.current_index]
                    
                    with st.spinner("🤖 AI is validating..."):
                        ai_result = validate_response_with_ai(user_input, current_placeholder, st.session_state.filled_data)
                    
                    if ai_result['should_proceed']:
                        st.session_state.filled_data[current_placeholder['key']] = user_input
                        st.session_state.current_index += 1
                        
                        if st.session_state.current_index < len(st.session_state.placeholders):
                            next_placeholder = st.session_state.placeholders[st.session_state.current_index]
                            next_question = get_ai_question_for_placeholder(
                                next_placeholder, 
                                st.session_state.filled_data,
                                st.session_state.document_text
                            )
                            response = f"{ai_result['response']} {next_question}"
                        else:
                            response = "Perfect! All placeholders have been filled. Your document is ready for review! 🎉"
                            st.session_state.step = 'complete'
                    else:
                        response = ai_result['response']
                    
                    st.session_state.messages.append({
                        'type': 'assistant',
                        'content': response,
                        'timestamp': datetime.now()
                    })
                    
                    st.rerun()
    
    with col2:
        st.markdown("### 📋 Placeholders")
        
        for idx, placeholder in enumerate(st.session_state.placeholders):
            if idx < st.session_state.current_index:
                status_class = "completed"
                icon = "✅"
            elif idx == st.session_state.current_index:
                status_class = "current"
                icon = "▶️"
            else:
                status_class = "pending"
                icon = "⭕"
            
            value_text = st.session_state.filled_data.get(placeholder['key'], '')
            value_display = f"<br><small><i>{value_text}</i></small>" if value_text else ""
            
            st.markdown(f"""
            <div class="placeholder-box {status_class}">
                {icon} <strong>{placeholder['label']}</strong>{value_display}
            </div>
            """, unsafe_allow_html=True)

elif st.session_state.step == 'complete':
    st.markdown("---")
    
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.markdown("""
        <div class="success-box">
            <h1>✅ Document Complete!</h1>
            <p>AI has generated your completed document. Review and download below.</p>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("<br>", unsafe_allow_html=True)
    
    with st.spinner("🤖 AI is generating your final document..."):
        completed_doc = generate_completed_document_with_ai()
        st.session_state.completed_doc = completed_doc
    
    st.markdown("### 📄 Document Preview")
    st.markdown(f'<div class="doc-preview">{completed_doc}</div>', unsafe_allow_html=True)
    
    st.markdown("<br>", unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 1, 1])
    
    with col1:
        st.download_button(
            label="📥 Download as TXT",
            data=completed_doc,
            file_name=f"completed_{st.session_state.file_name.replace('.docx', '.txt')}",
            mime="text/plain",
            use_container_width=True
        )
    
    with col2:
        doc = Document()
        for line in completed_doc.split('\n'):
            doc.add_paragraph(line)
        
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        
        st.download_button(
            label="📥 Download as DOCX",
            data=bio.getvalue(),
            file_name=f"completed_{st.session_state.file_name}",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    
    with col3:
        if st.button("🔄 Process Another", use_container_width=True):
            reset_app()

with st.sidebar:
    st.markdown("### 🤖 AI Status")
    if st.session_state.api_configured:
        st.success("✅ Gemini AI Connected")
    else:
        st.error("❌ AI Not Configured")
    
    st.markdown("---")
    
    st.markdown("### 📊 Statistics")
    if st.session_state.placeholders:
        st.metric("Total Placeholders", len(st.session_state.placeholders))
        st.metric("Filled", len(st.session_state.filled_data))
        st.metric("Remaining", len(st.session_state.placeholders) - len(st.session_state.filled_data))
    
    st.markdown("---")
    
    st.markdown("### ℹ️ About")
    st.info("""
    This app uses Gemini AI to:
    - Intelligently detect placeholders
    - Generate contextual questions
    - Validate your responses
    - Generate the final document
    
    All securely powered by your .env API key.
    """)
    
    st.markdown("---")
    st.caption("Built with Streamlit & Google Gemini AI")
    st.caption("🔒 API Key secured in .env file")
