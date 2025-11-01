import streamlit as st
import re
from docx import Document
from io import BytesIO
from datetime import datetime
# Note: Pillow (PIL) and lxml are imported by python-docx or used implicitly.
# If you ever need to use Pillow directly, you'd add: from PIL import Image

st.set_page_config(page_title="Document Placeholder Filler", page_icon="📝", layout="wide")

# --- All installation-related code has been removed ---

def extract_text_and_placeholders(docx_file):
    try:
        doc = Document(docx_file)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        text = "\n".join(full_text)
        patterns = [
            r'\{\{([^}]+)\}\}',
            r'\{([^}]+)\}',
            r'\[([^]]+)\]',
            r'<<([^>]+)>>',
        ]
        placeholders = set()
        for pattern in patterns:
            matches = re.findall(pattern, text)
            placeholders.update(matches)
        placeholders = [p.strip() for p in placeholders if len(p.strip()) > 1 and not p.strip().isdigit()]
        return text, sorted(list(set(placeholders)))
    except Exception as e:
        st.error(f"Error reading document: {str(e)}")
        return "", []

def generate_question_from_placeholder(placeholder, use_llm=False, api_key=None):
    if use_llm and api_key:
        try:
            # Your LLM logic would go here
            pass
        except Exception as e:
            st.warning(f"LLM generation failed: {str(e)}. Using fallback.")
    
    # Fallback readable question generation
    readable = placeholder.replace('_', ' ').replace('-', ' ')
    readable = re.sub(r'([a-z])([A-Z])', r'\1 \2', readable) # Add space for camelCase
    readable = readable.strip().capitalize()
    return f"What is the {readable}?"

def replace_placeholders_in_docx(original_file, placeholder_values):
    try:
        original_file.seek(0)
        doc = Document(original_file)
        
        # Process paragraphs
        for para in doc.paragraphs:
            for placeholder, value in placeholder_values.items():
                patterns_to_replace = [
                    f'{{{{{placeholder}}}}}',
                    f'{{{placeholder}}}',
                    f'[{placeholder}]',
                    f'<<{placeholder}>>',
                ]
                for pattern in patterns_to_replace:
                    if pattern in para.text:
                        # Need to iterate through runs to maintain formatting
                        for run in para.runs:
                            if pattern in run.text:
                                run.text = run.text.replace(pattern, value)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for placeholder, value in placeholder_values.items():
                        patterns_to_replace = [
                            f'{{{{{placeholder}}}}}',
                            f'{{{placeholder}}}',
                            f'[{placeholder}]',
                            f'<<{placeholder}>>',
                        ]
                        for pattern in patterns_to_replace:
                            if pattern in cell.text:
                                # Iterate through paragraphs in the cell
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        if pattern in run.text:
                                            run.text = run.text.replace(pattern, value)
        
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output
    except Exception as e:
        st.error(f"Error replacing placeholders: {str(e)}")
        return None

# --- Session State Initialization ---
if 'uploaded_file' not in st.session_state:
    st.session_state.uploaded_file = None
if 'placeholders' not in st.session_state:
    st.session_state.placeholders = []
if 'current_placeholder_index' not in st.session_state:
    st.session_state.current_placeholder_index = 0
if 'placeholder_values' not in st.session_state:
    st.session_state.placeholder_values = {}
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'document_completed' not in st.session_state:
    st.session_state.document_completed = False
if 'original_file_bytes' not in st.session_state:
    st.session_state.original_file_bytes = None

# --- Main App UI ---
st.title("📝 Document Placeholder Filler")
st.markdown("Upload a .docx document with placeholders, and I'll help you fill them in through a conversational interface!")

# --- Sidebar ---
with st.sidebar:
    st.header("⚙️ Settings")
    use_llm = st.checkbox("Use Gemini API for Questions", value=False, help="Generate more natural questions using AI")
    gemini_api_key = None
    if use_llm:
        gemini_api_key = st.text_input("Gemini API Key", type="password", help="Enter your Google Gemini API key")
    st.markdown("---")
    st.header("ℹ️ How it works")
    st.markdown("""
    1. Upload a `.docx` file with placeholders
    2. Supported formats:
        - `{{placeholder}}`
        - `{placeholder}`
        - `[placeholder]`
        - `<<placeholder>>`
    3. Answer questions in the chat
    4. Download your completed document
    """)
    st.markdown("---")
    if st.button("🔄 Start Over", use_container_width=True):
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.rerun()

# --- Main Page Logic ---

# 1. File Upload Section
if not st.session_state.document_completed:
    uploaded_file = st.file_uploader("Upload your .docx document", type=['docx'], help="Upload a Word document containing placeholders to fill")
    
    if uploaded_file is not None:
        # Check if it's a new file
        if st.session_state.uploaded_file is None or uploaded_file.name != st.session_state.uploaded_file:
            st.session_state.uploaded_file = uploaded_file.name
            st.session_state.original_file_bytes = uploaded_file.getvalue()
            
            # Process the new file
            file_bytes = BytesIO(st.session_state.original_file_bytes)
            text, placeholders = extract_text_and_placeholders(file_bytes)
            
            if placeholders:
                st.session_state.placeholders = placeholders
                st.session_state.current_placeholder_index = 0
                st.session_state.placeholder_values = {}
                st.session_state.chat_history = []
                st.session_state.document_completed = False
                st.success(f"✅ Document uploaded! Found {len(placeholders)} placeholders to fill.")
                with st.expander("📋 Detected Placeholders"):
                    for i, placeholder in enumerate(placeholders, 1):
                        st.write(f"{i}. `{placeholder}`")
            else:
                st.warning("⚠️ No placeholders detected in the document. Please ensure your document contains placeholders in supported formats.")
                st.session_state.uploaded_file = None # Reset to allow re-upload

# 2. Chat/Filling Section
if st.session_state.placeholders and not st.session_state.document_completed:
    st.markdown("---")
    st.header("💬 Fill in the Placeholders")
    
    chat_container = st.container()
    with chat_container:
        for message in st.session_state.chat_history:
            if message['role'] == 'assistant':
                with st.chat_message("assistant"):
                    st.write(message['content'])
            else:
                with st.chat_message("user"):
                    st.write(message['content'])

    # Check if we need to ask the next question
    if st.session_state.current_placeholder_index < len(st.session_state.placeholders):
        current_placeholder = st.session_state.placeholders[st.session_state.current_placeholder_index]
        
        # Display question if not already displayed
        if not st.session_state.chat_history or st.session_state.chat_history[-1]['role'] != 'assistant':
            question = generate_question_from_placeholder(
                current_placeholder,
                use_llm=use_llm,
                api_key=gemini_api_key if use_llm else None
            )
            st.session_state.chat_history.append({
                'role': 'assistant',
                'content': f"**Question {st.session_state.current_placeholder_index + 1}/{len(st.session_state.placeholders)}:** {question}"
            })
            st.rerun() # Rerun to display the new question

        # Get user input
        user_input = st.chat_input(f"Enter value for '{current_placeholder}'...")
        
        if user_input:
            st.session_state.placeholder_values[current_placeholder] = user_input
            st.session_state.chat_history.append({'role': 'user', 'content': user_input})
            st.session_state.current_placeholder_index += 1
            
            # Check if all placeholders are filled
            if st.session_state.current_placeholder_index >= len(st.session_state.placeholders):
                st.session_state.document_completed = True
                st.session_state.chat_history.append({'role': 'assistant', 'content': "🎉 Great! All placeholders have been filled. You can now download your completed document below."})
            
            st.rerun()

# 3. Download Section
if st.session_state.document_completed and st.session_state.original_file_bytes:
    st.markdown("---")
    st.header("✅ Document Complete!")
    
    original_file = BytesIO(st.session_state.original_file_bytes)
    completed_doc = replace_placeholders_in_docx(original_file, st.session_state.placeholder_values)
    
    if completed_doc:
        col1, col2 = st.columns(2)
        with col1:
            st.subheader("📄 Filled Values")
            for placeholder, value in st.session_state.placeholder_values.items():
                st.write(f"**{placeholder}:** {value}")
        
        with col2:
            st.subheader("⬇️ Download Options")
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            docx_filename = f"completed_document_{timestamp}.docx"
            
            st.download_button(
                label="📥 Download as DOCX",
                data=completed_doc.getvalue(),
                file_name=docx_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            st.info("💡 To convert to PDF, download the DOCX file and use Microsoft Word, Google Docs, or an online converter.")
        
        st.markdown("---")
        if st.button("✏️ Edit Values", use_container_width=True):
            st.session_state.document_completed = False
            st.session_state.current_placeholder_index = 0
            st.session_state.chat_history = []
            st.rerun()

# --- Footer ---
st.markdown("---")
st.markdown(
    """
    <div style='text-align: center; color: gray; font-size: 0.9em;'>
    Made By Sushant Kumar using Streamlit | Supports .docx files with multiple placeholder formats
    </div>
    """,
    unsafe_allow_html=True
)
